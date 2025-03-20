from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import seaborn as sns

class ReportFormatter:
    def __init__(self, doc):
        self.doc = doc

    def add_header(self, month, year):
        """Add styled header section"""
        header = self.doc.add_heading('Monthly Sales Report', 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add styled subtitle
        subtitle = self.doc.add_paragraph()
        subtitle_run = subtitle.add_run(f'Period: {month}/{year}')
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(70, 70, 70)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_metrics_table(self, metrics):
        """Add styled metrics table"""
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Colorful Grid Accent 1'
        
        for metric, value in metrics:
            row = table.add_row().cells
            row[0].text = metric
            row[1].text = value

    def add_product_analysis(self, products_data):
        """Add styled product analysis section"""
        section = self.doc.add_heading('Product Performance', level=1)
        section.style.font.color.rgb = RGBColor(0, 82, 136)
        
        for i, (product, data) in enumerate(products_data.iterrows(), 1):
            p = self.doc.add_paragraph(style='List Number')
            p.add_run(f'{product}\n').bold = True
            p.add_run(f'Revenue: ${data["Total Invoice Amount"]:,.2f}\n')
            p.add_run(f'Units Sold: {data["Quantity Sold"]:,}')

    def add_insights_section(self, insights):
        """Format AI insights in an attractive way"""
        section = self.doc.add_heading('AI-Generated Insights', level=1)
        section.style.font.color.rgb = RGBColor(0, 82, 136)
        
        # Convert markdown headers to styled paragraphs
        lines = insights.split('\n')
        current_paragraph = None
        
        for line in lines:
            if line.startswith('###'):
                # Subsection header
                p = self.doc.add_paragraph()
                run = p.add_run(line.replace('###', '').strip())
                run.bold = True
                run.font.size = Pt(12)
                current_paragraph = None
            elif line.startswith('-'):
                # Bullet point
                if not current_paragraph:
                    current_paragraph = self.doc.add_paragraph(style='List Bullet')
                current_paragraph.add_run(line.replace('-', '').strip())
                current_paragraph = None
            elif line.strip():
                # Regular text
                p = self.doc.add_paragraph()
                p.add_run(line.strip())

    def add_recommendations(self, recommendations):
        """Add styled recommendations section"""
        section = self.doc.add_heading('Strategic Recommendations', level=1)
        section.style.font.color.rgb = RGBColor(0, 82, 136)
        
        intro = self.doc.add_paragraph()
        intro.add_run('Based on our analysis, we recommend the following actions:').bold = True
        
        for rec in recommendations:
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(rec).font.size = Pt(11)

    def add_trend_chart(self, df_filtered, image_path):
        """Create and add styled trend chart"""
        plt.figure(figsize=(12, 6))
        # Remove seaborn style and use built-in style
        plt.style.use('bmh')  # Alternative built-in style
        
        # Create daily sales trend
        daily_sales = df_filtered.groupby(df_filtered['Invoice Date'].dt.day)['Total Invoice Amount'].sum()
        
        # Enhanced plotting with custom styling
        fig, ax = plt.subplots(figsize=(12, 6))
        
        # Plot line and points
        ax.plot(daily_sales.index, daily_sales.values, 
                color='#0066cc', 
                linewidth=2, 
                marker='o',
                markersize=6)
        
        # Add fill
        ax.fill_between(daily_sales.index, 
                        daily_sales.values, 
                        alpha=0.2, 
                        color='#0066cc')
        
        # Customize appearance
        ax.set_xlabel("Day of Month", fontsize=12, fontweight='bold')
        ax.set_ylabel("Revenue ($)", fontsize=12, fontweight='bold')
        ax.set_title("Daily Sales Trend", fontsize=14, pad=20, fontweight='bold')
        
        # Enhance grid
        ax.grid(True, linestyle='--', alpha=0.7)
        
        # Customize spines
        for spine in ax.spines.values():
            spine.set_color('#666666')
            spine.set_linewidth(0.5)
        
        # Format y-axis labels as currency
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x:,.0f}'))
        
        # Add padding
        plt.tight_layout()
        
        # Save with high quality
        plt.savefig(image_path, 
                    dpi=300, 
                    bbox_inches='tight',
                    facecolor='white',
                    edgecolor='none')
        plt.close()
        
        # Add to document
        self.doc.add_picture(image_path, width=Inches(6))
    
        # Add caption
        caption = self.doc.add_paragraph('Daily Sales Trend Analysis', style='Caption')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER