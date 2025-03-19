import logging
import json
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from phi.model.google import Gemini
from phi.agent.duckdb import DuckDbAgent
from docx import Document
from docx.shared import Inches
from datetime import datetime
from dotenv import load_dotenv
import os

load_dotenv()

logging.getLogger('phi').setLevel(logging.WARNING)

def generate_monthly_report(year, month):
    # Create reports directory
    reports_dir = "./reports"
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)
    
    # Generate filenames
    image_filename = f"sales_trend_{year}_{month}.png"
    image_path = os.path.join(reports_dir, image_filename)
    report_filename = f"sales_report_{year}_{month}.docx"
    report_path = os.path.join(reports_dir, report_filename)
    
    # Create sales trend graph
    df = pd.read_csv("./data/MOCK_DATA.csv")
    df['purchase_date'] = pd.to_datetime(df['purchase_date'])
    df_filtered = df[(df['purchase_date'].dt.year == year) & (df['purchase_date'].dt.month == month)]
    
    plt.figure(figsize=(10, 5))
    daily_sales = df_filtered.groupby(df_filtered['purchase_date'].dt.day)['total_price'].sum()
    sns.lineplot(x=daily_sales.index, y=daily_sales.values, marker='o')
    plt.xlabel("Day of the Month")
    plt.ylabel("Total Revenue ($)")
    plt.title(f"Daily Sales Trend - {month}/{year}")
    plt.grid()
    plt.savefig(image_path)
    plt.close()

    # Get sales data from DuckDB
    data_analyst_sales = DuckDbAgent(
        model=Gemini(model="gemini-2.0-flash"),
        semantic_model=json.dumps(
            {
                "tables": [
                    {
                        "name": "sales",
                        "description": "Contains retail sales transaction data including product details, pricing, payment methods and returns.",
                        "path": "./data/MOCK_DATA.csv",
                    }
                ]
            }
        ),
        markdown=True,
    )

    # Create Word document
    doc = Document()
    
    # Add header
    doc.add_heading(f'Monthly Sales Report - {month}/{year}', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Add Key Metrics section
    doc.add_heading('Key Metrics', level=1)
    
    # Get metrics using DuckDB queries
    total_revenue = df_filtered['total_price'].sum()
    num_transactions = len(df_filtered)
    avg_transaction = total_revenue / num_transactions
    return_rate = (df_filtered['is_returned'].value_counts().get(True, 0) / len(df_filtered)) * 100
    popular_payment = df_filtered['payment_method'].mode()[0]
    
    # Add metrics to document
    metrics = doc.add_paragraph()
    # Total Revenue
    run = metrics.add_run('Total Revenue: ')
    run.bold = True
    metrics.add_run(f'${total_revenue:,.2f}\n')
    
    # Number of Transactions
    run = metrics.add_run('Number of Transactions: ')
    run.bold = True
    metrics.add_run(f'{num_transactions:,}\n')
    
    # Average Transaction Value
    run = metrics.add_run('Average Transaction Value: ')
    run.bold = True
    metrics.add_run(f'${avg_transaction:,.2f}\n')
    
    # Return Rate
    run = metrics.add_run('Return Rate: ')
    run.bold = True
    metrics.add_run(f'{return_rate:.1f}%\n')
    
    # Most Popular Payment Method
    run = metrics.add_run('Most Popular Payment Method: ')
    run.bold = True
    metrics.add_run(f'{popular_payment}')

    # Add Top 5 Products section
    doc.add_heading('Top 5 Products by Revenue', level=2)
    top_products = df_filtered.groupby('product_name')['total_price'].sum().sort_values(ascending=False).head(5)
    
    for i, (product, revenue) in enumerate(top_products.items(), 1):
        doc.add_paragraph(f'{i}. {product} (${revenue:,.2f})', style='List Number')

    # Add Sales Trend section
    doc.add_heading('Sales Trends', level=1)
    doc.add_picture(image_path, width=Inches(6))
    
    # Save document
    doc.save(report_path)
    print(f"\nReport generated successfully: {report_path}")

    # Query AI for markdown report
    query = f"""Generate a sales report for {month}/{year} in this exact format:

    # Monthly Sales Report - {month}/{year}

    ## Key Metrics
    * Total Revenue: $X
    * Number of Transactions: X
    * Top 5 Products by Revenue:
        1. Product Name ($X)
        2. Product Name ($X)
        3. Product Name ($X)
        4. Product Name ($X)
        5. Product Name ($X)
    * Average Transaction Value: $X
    * Return Rate: X%
    * Most Popular Payment Method: X

    ## Sales Trends
    ![Daily Sales Trend](./reports/{image_filename})

    Note: Return just the markdown report with actual values, no explanations and no sql queries."""

    data_analyst_sales.print_response(query, stream=True)

if __name__ == "__main__":
    try:
        print("\nMonthly Sales Report Generator\n")
        year = int(input("Enter year (2021-2025): "))
        month = int(input("Enter month (1-12): "))
        
        if not (2021 <= year <= 2025) or not (1 <= month <= 12):
            print("\nError: Please enter a valid year (2021-2025) and month (1-12)")
        else:
            print("\nGenerating report...\n")
            generate_monthly_report(year, month)
    except ValueError:
        print("\nError: Please enter valid numeric values")