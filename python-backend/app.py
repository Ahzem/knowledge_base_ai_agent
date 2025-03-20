import logging
import json
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from datetime import datetime
from agents.sales_analyst import SalesAnalystAgent
from utils.report_formatter import ReportFormatter
from dotenv import load_dotenv
import os

load_dotenv()

logging.getLogger('phi').setLevel(logging.WARNING)

def get_ai_insights(df_filtered, month, year):
    """Get AI-generated insights about the sales data"""
    try:
        # Prepare data summary
        data_summary = {
            "period": {
                "month": month,
                "year": year
            },
            "financial_metrics": {
                "total_revenue": f"${df_filtered['Total Invoice Amount'].sum():,.2f}",
                "total_transactions": len(df_filtered),
                "avg_transaction": f"${df_filtered['Total Invoice Amount'].mean():,.2f}",
                "total_profit": f"${df_filtered['Gross Profit per Sale'].sum():,.2f}",
                "total_cogs": f"${df_filtered['Cost of Goods Sold (COGS)'].sum():,.2f}"
            },
            "product_metrics": {
                "top_products": df_filtered.groupby('Product Name').agg({
                    'Total Invoice Amount': 'sum',
                    'Quantity Sold': 'sum',
                    'Gross Profit per Sale': 'sum'
                }).sort_values('Total Invoice Amount', ascending=False).head(5).to_dict('index'),
                "category_performance": df_filtered.groupby('Product Category')['Total Invoice Amount'].sum().to_dict()
            },
            "customer_metrics": {
                "unique_customers": len(df_filtered['Customer ID'].unique()),
                "avg_customer_purchases": df_filtered.groupby('Customer ID').size().mean(),
                "top_customers": df_filtered.groupby('Customer Name')['Total Invoice Amount'].sum().sort_values(ascending=False).head(3).to_dict()
            },
            "discount_analysis": {
                "total_discounts": f"${df_filtered['Markdown/Discount Loss'].sum():,.2f}",
                "total_refunds": f"${df_filtered['Refund Amount'].sum():,.2f}",
                "avg_discount_rate": f"{df_filtered['Discount Applied'].mean():.1f}%"
            }
        }

        # Initialize and use the Sales Analyst Agent
        analyst = SalesAnalystAgent()
        return analyst.get_insights(json.dumps(data_summary, indent=2), month, year)

    except Exception as e:
        print(f"Error generating AI insights: {str(e)}")
        return "Error: Unable to generate insights due to technical issues."


def generate_monthly_report(year, month):
    # Create reports directory
    reports_dir = "./reports"
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)
    
    # Generate filenames
    image_filename = f"sales_trend_{year}_{month}.png"
    image_path = os.path.join(reports_dir, image_filename)
    report_filename = f"sales_report_{year}_{month}.docx"
    report_path = os.path.join(reports_dir, report_filename)
    
    # Create sales trend graph
    df = pd.read_csv("./data/FanBudget.csv")
    df['Invoice Date'] = pd.to_datetime(df['Invoice Date'])  # Changed from purchase_date
    df_filtered = df[(df['Invoice Date'].dt.year == year) & (df['Invoice Date'].dt.month == month)]
    
    plt.figure(figsize=(10, 5))
    daily_sales = df_filtered.groupby(df_filtered['Invoice Date'].dt.day)['Total Invoice Amount'].sum()  # Changed from total_price
    sns.lineplot(x=daily_sales.index, y=daily_sales.values, marker='o')
    plt.xlabel("Day of the Month")
    plt.ylabel("Total Revenue ($)")
    plt.title(f"Daily Sales Trend - {month}/{year}")
    plt.grid()
    plt.savefig(image_path)
    plt.close()

    # Create Word document with enhanced formatting
    doc = Document()
    formatter = ReportFormatter(doc)
    
    # Add styled header
    formatter.add_header(month, year)
    
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('=' * 60)

    # Enhanced metrics calculations
    total_revenue = df_filtered['Total Invoice Amount'].sum()
    num_transactions = len(df_filtered)
    avg_transaction = total_revenue / num_transactions if num_transactions > 0 else 0
    return_rate = (df_filtered['Refund Amount'] > 0).mean() * 100
    
    # Calculate additional metrics
    daily_avg_revenue = total_revenue / df_filtered['Invoice Date'].nunique()
    customer_counts = df_filtered['Customer ID'].value_counts()
    repeat_rate = (len(customer_counts[customer_counts > 1]) / len(df_filtered['Customer ID'].unique())) * 100 if len(df_filtered) > 0 else 0

    # Add Key Metrics section with enhanced formatting
    doc.add_heading('Key Performance Metrics', level=1)
    
        
    metrics_table = doc.add_table(rows=1, cols=2)
    metrics_table.style = 'Table Grid'
    
    # Add metrics in a structured table
    metrics = [
        ('Total Revenue', f'${total_revenue:,.2f}'),
        ('Number of Transactions', f'{num_transactions:,}'),
        ('Average Transaction Value', f'${avg_transaction:,.2f}'),
        ('Daily Average Revenue', f'${daily_avg_revenue:,.2f}'),
        ('Return Rate', f'{return_rate:.1f}%'),
        ('Customer Repeat Rate', f'{repeat_rate:.1f}%')
    ]
    
    formatter.add_metrics_table(metrics)

    # Add Product Analysis section
    doc.add_heading('Product Performance Analysis', level=1)
    
    # Top products analysis
    top_products = df_filtered.groupby('Product Name').agg({
        'Total Invoice Amount': 'sum',
        'Quantity Sold': 'sum'
    }).sort_values('Total Invoice Amount', ascending=False).head(5)
    
    formatter.add_product_analysis(top_products)

    # Add Sales Trend Analysis
    formatter.add_trend_chart(df_filtered, image_path)
    
    # Get AI Insights
    insights = get_ai_insights(df_filtered, month, year)
    formatter.add_insights_section(insights)

    
    # Add recommendations
    recommendations = [
        "Optimize inventory levels for top-selling products to prevent stockouts",
        f"Focus on customer retention strategies given the {repeat_rate:.1f}% repeat rate",
        f"Review pricing strategy for products with high return rates ({return_rate:.1f}%)",
        "Implement targeted promotions based on customer purchase patterns",
        "Consider expanding product categories with consistent performance"
    ]
    formatter.add_recommendations(recommendations)
    
    # Save document
    doc.save(report_path)
    print(f"\nEnhanced report generated successfully: {report_path}")

if __name__ == "__main__":
    try:
        print("\nMonthly Sales Report Generator\n")
        year = int(input("Enter year (2021-2025): "))
        month = int(input("Enter month (1-12): "))
        
        if not (2021 <= year <= 2025) or not (1 <= month <= 12):
            print("\nError: Please enter a valid year (2021-2025) and month (1-12)")
        else:
            print("\nGenerating report...\n")
            generate_monthly_report(year, month)
    except ValueError:
        print("\nError: Please enter valid numeric values")