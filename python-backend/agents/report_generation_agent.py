from phi.agent import Agent
from phi.model.openai import OpenAIChat
from phi.tools.pandas import PandasTools
from .financial_analytics_agent import FinancialAnalyticsAgent
from .recommendation_agent import RecommendationAgent
from .visualization_agent import VisualizationAgent
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import datetime
import os
import calendar
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')

class ReportGenerationAgent:
    def __init__(self):
        # Initialize component agents
        self.financial_agent = FinancialAnalyticsAgent()
        self.recommendation_agent = RecommendationAgent()
        self.visualization_agent = VisualizationAgent()
        
        # Initialize the report generation agent
        self.agent = Agent(
            name="Creative Report Generation Expert",
            model=OpenAIChat(
                model="gpt-4o",
                temperature=0.7,
                system=self.get_system_prompt()
            ),
            tools=[
                PandasTools(),
                self._get_financial_metrics,
                self._get_product_metrics,
                self._get_customer_metrics,
                self._add_executive_insights,
                self._get_previous_period_data
            ],
            description="Expert in creating comprehensive, visually engaging business reports",
            instructions=[
                "Create compelling narrative connecting data points",
                "Find creative ways to present standard business metrics",
                "Include visual metaphors and clear explanations",
                "Present actionable insights prominently",
                "Ensure consistency in tone and design throughout"
            ],
            add_history_to_messages=True,
            num_history_responses=5,
            show_tool_calls=False,
            markdown=True,
        )
        
        # Report design settings
        self.color_scheme = {
            'primary': RGBColor(41, 71, 135),    # Deep blue
            'secondary': RGBColor(233, 131, 0),  # Orange
            'positive': RGBColor(56, 142, 60),   # Green
            'negative': RGBColor(211, 47, 47),   # Red
            'neutral': RGBColor(117, 117, 117)   # Gray
        }
        
        # Report template settings
        self.report_templates = {
            "monthly": {
                "sections": [
                    "Executive Summary",
                    "Monthly Performance Highlights",
                    "Financial Analysis",
                    "Product Performance",
                    "Customer Insights",
                    "Recommendations",
                    "Next Month Outlook"
                ],
                "charts": ["revenue_trend", "product_comparison", "customer_metrics"]
            },
            "yearly": {
                "sections": [
                    "Executive Summary",
                    "Year in Review",
                    "Financial Performance",
                    "Product Portfolio Analysis",
                    "Customer Base Evolution",
                    "Market Position",
                    "Strategic Recommendations",
                    "Next Year Forecast"
                ],
                "charts": ["yearly_performance", "quarterly_comparison", "category_evolution", "customer_segments"]
            }
        }

    def _get_financial_metrics(self):
        """Tool to get financial metrics from DataSummaryAgent"""
        return self.financial_agent.data_agent.data_summary.get("financial_metrics", {})
    
    def _get_product_metrics(self):
        """Tool to get product metrics from DataSummaryAgent"""
        return self.financial_agent.data_agent.data_summary.get("product_metrics", {})
    
    def _get_customer_metrics(self):
        """Tool to get customer metrics from DataSummaryAgent"""
        return self.financial_agent.data_agent.data_summary.get("customer_metrics", {})
    
    def _add_executive_insights(self, query):
        """Tool to generate executive insights based on data"""
        try:
            # Use financial agent for insightful analysis
            insights = self.financial_agent.analyze(
                f"Provide 3-5 key executive insights for a business report: {query}"
            )
            return {"insights": insights}
        except Exception as e:
            return {"error": str(e)}
    
    def _get_previous_period_data(self, period_type, current_period):
        """Tool to get comparative data from previous period"""
        try:
            if period_type == "monthly":
                year, month = current_period
                # Calculate previous month
                prev_month = month - 1
                prev_year = year
                if prev_month < 1:
                    prev_month = 12
                    prev_year -= 1
                previous_period = (prev_year, prev_month)
            elif period_type == "yearly":
                year = current_period[0]
                previous_period = (year - 1,)
            else:
                return {"error": "Invalid period type"}
                
            # Get comparison data
            comparison = self.financial_agent.get_period_comparison(
                current_period=current_period,
                previous_period=previous_period
            )
            
            return {"comparison": comparison}
        except Exception as e:
            return {"error": str(e)}

    @staticmethod
    def get_system_prompt():
        return """You are an expert report generation system specializing in creating vibrant, insightful business reports that blend data analysis with compelling narratives. For all reports:

            1. Report Structure and Style:
            - Begin with an engaging, visually distinctive Executive Summary
            - Use a storytelling approach that connects data points into a coherent business narrative
            - Create clear section transitions that maintain reader interest
            - Include data visualizations with explanatory captions
            - Highlight key metrics with contextual explanations of their significance
            - Use metaphors and analogies to make complex financial concepts accessible

            2. Content Requirements:
            - Transform data into actionable business intelligence
            - Include comparative analysis (month-over-month or year-over-year)
            - Identify emerging patterns and anomalies in the data
            - Provide forward-looking projections with confidence levels
            - Include specific action recommendations with expected outcomes
            - Address both opportunities and challenges revealed by the data

            3. Creative Presentation:
            - Use storytelling techniques to maintain reader engagement
            - Incorporate visual metaphors that explain business trends
            - Create memorable section headings that convey key themes
            - Include pull quotes highlighting critical insights
            - Use color-coding to indicate performance (green for positive, red for concerning)
            - Create a consistent visual language throughout the report

            Remember that your reports should be both analytically rigorous and creatively engaging, striking a balance between data-driven insights and accessible business narrative.
            """

    def _add_future_outlook(self, doc, predictive_insights, report_type, period):
        """Add future outlook section with predictions"""
        doc.add_heading("Future Outlook", 1)
        
        # Add introduction based on report type
        intro_para = doc.add_paragraph()
        if report_type == "monthly":
            year, month = period
            month_name = calendar.month_name[month]
            next_month = month + 1 if month < 12 else 1
            next_month_name = calendar.month_name[next_month]
            next_year = year if month < 12 else year + 1
            intro_para.add_run(
                f"Based on current trends and predictive analysis, this section outlines "
                f"expected performance for {next_month_name} {next_year} and beyond."
            )
        else:
            year = period[0]
            intro_para.add_run(
                f"Based on annual performance analysis and market trends, this section "
                f"outlines projected performance for {year + 1}."
            )
        
        # Process predictive insights
        if predictive_insights:
            # Add predictions section
            doc.add_heading("Key Predictions", 2)
            
            # Parse and format predictions
            predictions = predictive_insights.split('\n\n')
            for prediction in predictions:
                if prediction.strip():
                    para = doc.add_paragraph()
                    para.add_run(prediction.strip())
            
            # Add risk factors
            doc.add_heading("Risk Factors", 2)
            risk_para = doc.add_paragraph()
            risk_para.add_run(
                "The following factors could impact our projections and should be monitored:"
            )
            
            # Common risk factors
            risks = [
                ("Market Volatility", "Economic conditions and market changes"),
                ("Customer Behavior", "Shifts in purchasing patterns or preferences"),
                ("Competition", "New market entrants or competitive actions"),
                ("Supply Chain", "Potential disruptions or cost variations"),
                ("Operational", "Internal process changes and adaptations")
            ]
            
            # Add risks as bullet points
            for risk_type, description in risks:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.add_run(f"{risk_type}: ").bold = True
                bullet.add_run(description)
        
        else:
            # If no predictive insights available
            no_data_para = doc.add_paragraph()
            no_data_para.add_run(
                "Detailed predictive insights are not available for this period. "
                "Please ensure all required data is present for future forecasting."
            ).italic = True
        
        # Add page break after future outlook
        doc.add_page_break()

    def _add_appendix(self, doc):
        """Add appendix with methodology information"""
        doc.add_heading("Appendix", 1)
        
        # Add methodology section
        doc.add_heading("Methodology", 2)
        methodology_para = doc.add_paragraph()
        methodology_para.add_run(
            "This report was generated using advanced data analytics and business intelligence "
            "tools. The analysis includes:"
        )
        
        # Add methodology details
        methodologies = [
            "Historical data analysis and trend identification",
            "Comparative period-over-period metrics",
            "Statistical analysis of key performance indicators",
            "Predictive modeling for future projections",
            "Industry standard financial calculations and ratios"
        ]
        
        for method in methodologies:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(method)
        
        # Add data sources
        doc.add_heading("Data Sources", 2)
        sources_para = doc.add_paragraph()
        sources_para.add_run(
            "The analysis in this report is based on data from the following sources:"
        )
        
        # Add source details
        sources = [
            "Transaction and sales records",
            "Customer relationship management system",
            "Inventory management system",
            "Financial accounting records",
            "Market research and industry data"
        ]
        
        for source in sources:
            bullet = doc.add_paragraph(style='List Bullet')
            bullet.add_run(source)
            
    def generate_report(self, report_type="monthly", period=None, title=None, output_path=None):
        """Generate a complete business report document"""
        try:
            # Get current period if not specified
            if not period:
                now = datetime.now()
                if report_type == "monthly":
                    period = (now.year, now.month)
                else:
                    period = (now.year,)
            
            # Generate default title if not provided
            if not title:
                if report_type == "monthly":
                    month_name = calendar.month_name[period[1]]
                    title = f"{month_name} {period[0]} Business Performance Report"
                else:
                    title = f"{period[0]} Annual Business Performance Review"
            
            # Generate default output path if not provided
            if not output_path:
                os.makedirs('reports', exist_ok=True)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                if report_type == "monthly":
                    output_path = f"reports/monthly_report_{period[0]}_{period[1]}_{timestamp}.docx"
                else:
                    output_path = f"reports/annual_report_{period[0]}_{timestamp}.docx"
            
            # Create a new document
            doc = Document()
            
            # Set up document properties and styles
            self._setup_document_properties(doc, report_type, period, title)
            
            # Add cover page
            self._add_cover_page(doc, report_type, period, title)
            
            # Add table of contents
            self._add_table_of_contents(doc)
            
            # Add executive summary
            self._add_executive_summary(doc, report_type, period)
            
            # Add financial analysis
            self._add_financial_analysis(doc, report_type, period)
            
            # Add product analysis
            self._add_product_analysis(doc, report_type, period)
            
            # Add customer analysis
            self._add_customer_analysis(doc, report_type, period)
            
            # Add recommendations
            self._add_recommendations(doc, report_type, period)
            
            # Add future outlook
            insights = self.recommendation_agent.get_recommendations(
                f"Provide predictive insights for the next period after {period}"
            )
            self._add_future_outlook(doc, insights, report_type, period)
            
            # Add appendix
            self._add_appendix(doc)
            
            # Save the document
            doc.save(output_path)
            
            print(f"Report generated successfully: {output_path}")
            return output_path
            
        except Exception as e:
            error_msg = f"Error generating report: {str(e)}"
            print(error_msg)
            return error_msg
            
        except Exception as e:
            print(f"Error generating report: {str(e)}")
            raise
    
    def _setup_document_properties(self, doc, report_type, period, title=None):
        """Set up document properties and styles"""
        # Set default document styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add custom heading styles
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = 'Calibri'
            heading_font.bold = True
            heading_font.color.rgb = self.color_scheme['primary']
        
        # Set document properties
        if title:
            doc.core_properties.title = title
        else:
            if report_type == "monthly":
                year, month = period
                month_name = calendar.month_name[month]
                doc.core_properties.title = f"Monthly Business Report: {month_name} {year}"
            else:
                year = period[0]
                doc.core_properties.title = f"Annual Business Report: {year}"
        
        doc.core_properties.author = "Business Intelligence System"
        doc.core_properties.created = datetime.now()
    
    def _add_cover_page(self, doc, report_type, period, title=None):
        """Add a professional cover page to document"""
        # Add section with different orientation if needed
        section = doc.sections[0]
        
        # Add title
        if title:
            report_title = title
        else:
            if report_type == "monthly":
                year, month = period
                month_name = calendar.month_name[month]
                report_title = f"Monthly Business Report\n{month_name} {year}"
            else:
                year = period[0]
                report_title = f"Annual Business Report\n{year}"
        
        # Add large title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(report_title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = self.color_scheme['primary']
        
        # Add space
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add company logo if available
        logo_path = "assets/company_logo.png"
        if os.path.exists(logo_path):
            doc.add_picture(logo_path, width=Inches(3))
            logo_para = doc.paragraphs[-1]
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add space
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.size = Pt(12)
        date_run.font.italic = True
        
        # Add company name
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run("Efito Solution (Pvt) Ltd")
        company_run.font.size = Pt(14)
        company_run.font.bold = True
        
        # Add page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc):
        """Add table of contents to document"""
        doc.add_heading("Table of Contents", 1)
        
        # Add placeholder for TOC (would need to be updated manually in Word)
        toc_para = doc.add_paragraph()
        toc_para.add_run("Right-click here and select 'Update Field' to update the table of contents.")
        
        # XML for TOC field
        toc_xml = '''
        <w:sdt>
            <w:sdtPr>
                <w:docPartObj>
                    <w:docPartGallery w:val="Table of Contents"/>
                    <w:docPartUnique/>
                </w:docPartObj>
            </w:sdtPr>
            <w:sdtContent>
                <w:p>
                    <w:r>
                        <w:fldChar w:fldCharType="begin"/>
                    </w:r>
                    <w:r>
                        <w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>
                    </w:r>
                    <w:r>
                        <w:fldChar w:fldCharType="separate"/>
                    </w:r>
                    <w:r>
                        <w:t>Table of Contents placeholder - Update in Word</w:t>
                    </w:r>
                    <w:r>
                        <w:fldChar w:fldCharType="end"/>
                    </w:r>
                </w:p>
            </w:sdtContent>
        </w:sdt>
        '''
        
        # Add page break
        doc.add_page_break()
    
    def _add_executive_summary(self, doc, report_type, period):
        """Add executive summary section"""
        doc.add_heading("Executive Summary", 1)
        
        # Get period info for context
        if report_type == "monthly":
            year, month = period
            month_name = calendar.month_name[month]
            period_str = f"{month_name} {year}"
        else:
            year = period[0]
            period_str = f"Year {year}"
        
        # Add introduction paragraph
        intro_para = doc.add_paragraph()
        intro_para.add_run(
            f"This report provides a comprehensive analysis of business performance for {period_str}. "
            f"It examines financial metrics, product performance, customer behavior, and offers "
            f"strategic recommendations based on data-driven insights."
        )
        
        # Get executive insights
        insights_query = f"Provide key executive insights for {period_str} based on our financial and customer data"
        insights_response = self._add_executive_insights(insights_query)
        executive_insights = insights_response.get("insights", "No insights available")
        
        # Add insights
        doc.add_heading("Key Insights", 2)
        insights_para = doc.add_paragraph()
        insights_para.add_run(executive_insights)
        
        # Add financial highlights
        doc.add_heading("Financial Highlights", 2)
        
        # Get financial metrics
        metrics = self.financial_agent.data_agent.data_summary.get("financial_metrics", {})
        
        # Create a table for key metrics
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Metric"
        header_cells[1].text = "Value"
        
        # Add key metrics
        key_metrics = [
            ("Total Revenue", metrics.get("total_revenue_formatted", "N/A")),
            ("Total Profit", metrics.get("total_profit_formatted", "N/A")),
            ("Average Transaction", metrics.get("avg_transaction_formatted", "N/A")),
            ("Total Transactions", str(metrics.get("total_transactions", "N/A")))
        ]
        
        for metric, value in key_metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # Add space after table
        doc.add_paragraph()
        
        # Add page break after executive summary
        doc.add_page_break()
    
    def _add_financial_analysis(self, doc, report_type, period):
        """Add financial analysis section"""
        doc.add_heading("Financial Analysis", 1)
        
        # Get period info for context
        if report_type == "monthly":
            year, month = period
            month_name = calendar.month_name[month]
            period_str = f"{month_name} {year}"
        else:
            year = period[0]
            period_str = f"Year {year}"
        
        # Add introduction paragraph
        intro_para = doc.add_paragraph()
        intro_para.add_run(
            f"This section presents a detailed analysis of financial performance for {period_str}, "
            f"including revenue, profitability, and key financial indicators."
        )
        
        # Get financial metrics and comparisons
        metrics = self.financial_agent.data_agent.data_summary.get("financial_metrics", {})
        comparison_data = self._get_previous_period_data(report_type, period)
        comparison = comparison_data.get("comparison", {})
        
        # Add revenue analysis
        doc.add_heading("Revenue Analysis", 2)
        revenue_para = doc.add_paragraph()
        
        # Check if we have comparison data
        if comparison and "financial" in comparison:
            current_revenue = metrics.get("total_revenue", 0)
            prev_revenue = comparison.get("financial", {}).get("total_revenue", {}).get("previous", 0)
            revenue_change = current_revenue - prev_revenue
            revenue_pct = (revenue_change / prev_revenue * 100) if prev_revenue != 0 else 0
            
            # Format revenue trend
            if revenue_change > 0:
                trend_color = self.color_scheme['positive']
                trend_text = f"an increase of {abs(revenue_pct):.1f}%"
            elif revenue_change < 0:
                trend_color = self.color_scheme['negative']
                trend_text = f"a decrease of {abs(revenue_pct):.1f}%"
            else:
                trend_color = self.color_scheme['neutral']
                trend_text = "no change"
                
            # Add revenue paragraph with colored trend
            revenue_para.add_run(
                f"Total revenue for {period_str} was {metrics.get('total_revenue_formatted', '$0.00')}, "
            )
            trend_run = revenue_para.add_run(trend_text)
            trend_run.font.color.rgb = trend_color
            revenue_para.add_run(
                f" compared to the previous period. The daily average revenue was "
                f"{metrics.get('daily_avg_revenue_formatted', '$0.00')}."
            )
        else:
            # No comparison data
            revenue_para.add_run(
                f"Total revenue for {period_str} was {metrics.get('total_revenue_formatted', '$0.00')}. "
                f"The daily average revenue was {metrics.get('daily_avg_revenue_formatted', '$0.00')}."
            )
        
        # Add profitability analysis
        doc.add_heading("Profitability Analysis", 2)
        profit_para = doc.add_paragraph()
        
        if "total_profit" in metrics and "total_revenue" in metrics:
            profit_margin = (metrics["total_profit"] / metrics["total_revenue"] * 100) if metrics["total_revenue"] > 0 else 0
            profit_para.add_run(
                f"Total profit for {period_str} was {metrics.get('total_profit_formatted', '$0.00')}, "
                f"representing a profit margin of {profit_margin:.1f}%. "
            )
            
            # Add COGS information
            profit_para.add_run(
                f"The total cost of goods sold was {metrics.get('total_cogs_formatted', '$0.00')}."
            )
        else:
            profit_para.add_run("Profit data is not available for this period.")
        
        # Add chart
        self._add_financial_chart(doc, period_str)
        
        # Add page break
        doc.add_page_break()
    
    def _add_financial_chart(self, doc, period_str):
        """Add financial chart to document"""
        # Create a simple bar chart using matplotlib
        try:
            # Get metrics
            metrics = self.financial_agent.data_agent.data_summary.get("financial_metrics", {})
            
            # Create data for chart
            labels = ['Revenue', 'Profit', 'COGS']
            values = [
                metrics.get('total_revenue', 0),
                metrics.get('total_profit', 0),
                metrics.get('total_cogs', 0)
            ]
            
            # Create the chart
            fig, ax = plt.figure(figsize=(7, 4)), plt.axes()
            bars = ax.bar(labels, values, color=['#295187', '#38903C', '#D52F2F'])
            
            # Add value labels on top of bars
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                        f'${height:,.0f}', ha='center', va='bottom', rotation=0)
            
            # Set title and labels
            ax.set_title(f'Financial Summary: {period_str}')
            ax.set_ylabel('Amount ($)')
            
            # Save chart to temporary file
            os.makedirs('temp', exist_ok=True)
            chart_path = 'temp/financial_chart.png'
            plt.tight_layout()
            plt.savefig(chart_path)
            plt.close()
            
            # Add chart to document
            doc.add_picture(chart_path, width=Inches(6))
            chart_para = doc.paragraphs[-1]
            chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add caption
            caption_para = doc.add_paragraph("Figure: Financial performance summary showing revenue, profit, and cost of goods sold.")
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.style = 'Caption'
            
            # Clean up
            if os.path.exists(chart_path):
                os.remove(chart_path)
                
        except Exception as e:
            # If chart creation fails, add a placeholder paragraph
            error_para = doc.add_paragraph()
            error_para.add_run(f"Chart could not be generated: {str(e)}").italic = True
    
    def _add_product_analysis(self, doc, report_type, period):
        """Add product performance analysis section"""
        doc.add_heading("Product Performance", 1)
        
        # Get period info for context
        if report_type == "monthly":
            year, month = period
            month_name = calendar.month_name[month]
            period_str = f"{month_name} {year}"
        else:
            year = period[0]
            period_str = f"Year {year}"
        
        # Add introduction paragraph
        intro_para = doc.add_paragraph()
        intro_para.add_run(
            f"This section analyzes product performance for {period_str}, identifying top performers, "
            f"trends, and opportunities for improvement."
        )
        
        # Get product metrics
        product_metrics = self.financial_agent.data_agent.data_summary.get("product_metrics", {})
        top_products = product_metrics.get("top_products", {})
        
        # Add top products analysis
        doc.add_heading("Top Performing Products", 2)
        
        if top_products:
            # Create table for top products
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            header_cells[0].text = "Product Name"
            header_cells[1].text = "Revenue"
            header_cells[2].text = "Units Sold"
            
            # Add product data
            for product_name, data in top_products.items():
                row_cells = table.add_row().cells
                row_cells[0].text = product_name
                row_cells[1].text = data.get("Total Invoice Amount_formatted", "N/A") 
                row_cells[2].text = str(int(data.get("Quantity Sold", 0)))
        else:
            no_data_para = doc.add_paragraph()
            no_data_para.add_run("No product performance data available for this period.").italic = True
        
        # Add category performance
        doc.add_heading("Category Performance", 2)
        
        category_data = product_metrics.get("category_performance", {})
        if category_data:
            # Create table for categories
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            header_cells[0].text = "Category"
            header_cells[1].text = "Revenue"
            
            # Add category data
            for category, data in category_data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = category
                row_cells[1].text = data.get("revenue_formatted", "N/A")
        else:
            no_data_para = doc.add_paragraph()
            no_data_para.add_run("No category performance data available for this period.").italic = True
        
        # Add product chart
        self._add_product_chart(doc, period_str)
        
        # Add page break
        doc.add_page_break()
    
    def _add_product_chart(self, doc, period_str):
        """Add product performance chart to document"""
        try:
            # Get product metrics
            product_metrics = self.financial_agent.data_agent.data_summary.get("product_metrics", {})
            top_products = product_metrics.get("top_products", {})
            
            if top_products:
                # Extract data for chart
                products = list(top_products.keys())[:5]  # Limit to top 5
                revenues = [data.get("Total Invoice Amount", 0) for data in list(top_products.values())[:5]]
                
                # Create the chart
                fig, ax = plt.figure(figsize=(7, 4)), plt.axes()
                bars = ax.bar(products, revenues, color='#295187')
                
                # Add value labels on top of bars
                for bar in bars:
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                            f'${height:,.0f}', ha='center', va='bottom', rotation=0)
                
                # Set title and labels
                ax.set_title(f'Top Products by Revenue: {period_str}')
                ax.set_ylabel('Revenue ($)')
                plt.xticks(rotation=45, ha='right')
                
                # Save chart to temporary file
                os.makedirs('temp', exist_ok=True)
                chart_path = 'temp/product_chart.png'
                plt.tight_layout()
                plt.savefig(chart_path)
                plt.close()
                
                # Add chart to document
                doc.add_picture(chart_path, width=Inches(6))
                chart_para = doc.paragraphs[-1]
                chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add caption
                caption_para = doc.add_paragraph("Figure: Revenue performance of top products.")
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.style = 'Caption'
                
                # Clean up
                if os.path.exists(chart_path):
                    os.remove(chart_path)
            else:
                error_para = doc.add_paragraph()
                error_para.add_run("Product chart could not be generated: No product data available").italic = True
                
        except Exception as e:
            # If chart creation fails, add a placeholder paragraph
            error_para = doc.add_paragraph()
            error_para.add_run(f"Product chart could not be generated: {str(e)}").italic = True
    
    def _add_customer_analysis(self, doc, report_type, period):
        """Add customer analysis section"""
        doc.add_heading("Customer Insights", 1)
        
        # Get period info for context
        if report_type == "monthly":
            year, month = period
            month_name = calendar.month_name[month]
            period_str = f"{month_name} {year}"
        else:
            year = period[0]
            period_str = f"Year {year}"
        
        # Add introduction paragraph
        intro_para = doc.add_paragraph()
        intro_para.add_run(
            f"This section examines customer behavior and patterns for {period_str}, "
            f"including customer acquisition, retention, and spending habits."
        )
        
        # Get customer metrics
        customer_metrics = self.financial_agent.data_agent.data_summary.get("customer_metrics", {})
        
        # Add customer overview
        doc.add_heading("Customer Overview", 2)
        overview_para = doc.add_paragraph()
        
        unique_customers = customer_metrics.get("unique_customers", 0)
        avg_purchases = customer_metrics.get("avg_customer_purchases", 0)
        repeat_rate = customer_metrics.get("repeat_rate", 0)
        
        overview_para.add_run(
            f"During {period_str}, we served {unique_customers} unique customers with an average of "
            f"{avg_purchases:.2f} purchases per customer. Our customer repeat rate was "
            f"{customer_metrics.get('repeat_rate_formatted', '0.0%')}."
        )
        
        # Add top customers section
        doc.add_heading("Top Customers", 2)
        
        top_customers = customer_metrics.get("top_customers", {})
        if top_customers:
            # Create table for top customers
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            header_cells[0].text = "Customer Name"
            header_cells[1].text = "Total Spend"
            
            # Add customer data
            for customer, data in top_customers.items():
                row_cells = table.add_row().cells
                row_cells[0].text = customer
                row_cells[1].text = data.get("spend_formatted", "N/A")
        else:
            no_data_para = doc.add_paragraph()
            no_data_para.add_run("No top customer data available for this period.").italic = True
        
        # Add page break
        doc.add_page_break()
    
    def _add_recommendations(self, doc, report_type, period):
        """Add strategic recommendations section"""
        doc.add_heading("Strategic Recommendations", 1)
        
        # Get period info for context
        if report_type == "monthly":
            year, month = period
            month_name = calendar.month_name[month]
            period_str = f"{month_name} {year}"
        else:
            year = period[0]
            period_str = f"Year {year}"
        
        # Add introduction paragraph
        intro_para = doc.add_paragraph()
        intro_para.add_run(
            f"Based on the data analysis for {period_str}, we recommend the following "
            f"strategic actions to improve business performance."
        )
        
        # Get recommendations from the agent
        recommendations = self.recommendation_agent.get_recommendations(
            f"Generate strategic business recommendations based on data for {period_str}",
            context=f"For {report_type} report covering {period_str}"
        )
        
        # Add recommendations to document
        if recommendations:
            # Split recommendations into paragraphs and add to document
            recommendation_sections = recommendations.split("\n\n")
            for section in recommendation_sections:
                if section.strip():
                    # Check if this looks like a heading
                    if section.startswith('#'):
                        # Process markdown heading
                        lines = section.split("\n")
                        heading = lines[0].lstrip('#').strip()
                        heading_level = min(len(lines[0]) - len(lines[0].lstrip('#')) + 1, 3)
                        doc.add_heading(heading, heading_level)
                        
                        # Process remaining lines in this section
                        if len(lines) > 1:
                            content = "\n".join(lines[1:])
                            if content.strip():
                                para = doc.add_paragraph()
                                para.add_run(content.strip())
                    else:
                        # Regular paragraph
                        para = doc.add_paragraph()
                        para.add_run(section.strip())
        else:
            no_recs_para = doc.add_paragraph()
            no_recs_para.add_run(
                "Unable to generate recommendations for this period. Please ensure "
                "all required data is available."
            ).italic = True
        
        # Add page break
        doc.add_page_break()

# Example usage
if __name__ == "__main__":
    report_generator = ReportGenerationAgent()
    
    # Generate a monthly report
    monthly_report_path = report_generator.generate_report(
        report_type="monthly",
        period=(2024, 3),  # March 2024
        title="Monthly Business Performance Report"
    )
    
    print(f"Monthly report generated: {monthly_report_path}")
    
    # Generate an annual report
    annual_report_path = report_generator.generate_report(
        report_type="yearly",
        period=(2023,),  # Year 2023
        title="Annual Business Performance Review"
    )
    
    print(f"Annual report generated: {annual_report_path}")